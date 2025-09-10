import re
import datetime
import pathlib
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from zipfile import BadZipFile
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter
import sys                                 # 打包兼容
import os
import tempfile
import time
import shutil

# ================== 0. 路径 & 常量 ==================
def resource_path(rel: str | pathlib.Path) -> pathlib.Path:
    """打包前/后均可用的资源绝对路径"""
    if getattr(sys, "frozen", False):           # EXE 中
        base = pathlib.Path(sys._MEIPASS)
    else:                                       # 源代码
        base = pathlib.Path(__file__).resolve().parent
    return base / rel

# ➤ exe/脚本所在目录
if getattr(sys, "frozen", False):
    APP_DIR = pathlib.Path(sys.executable).resolve().parent
else:
    APP_DIR = pathlib.Path(__file__).resolve().parent

# ➤ 所有输出集中到这里
OUTPUT_DIR = APP_DIR / "输出文件"
OUTPUT_DIR.mkdir(exist_ok=True)

TEMPLATE_DIR = resource_path("走读式模板V2.2")

PLACEHOLDER_RE = re.compile(r"`([^`]+)`")

SPECIAL_TEMPLATE_STEM = "A02办案安全承诺书"
B05_TEMPLATE_STEM = "B05安全评估表"  # ★ B05 模板名（不含扩展名）
ROLE_PHRASE = "(办案组组长、办案人员、安全员)"
ROLES = ["办案组组长", "办案人员", "安全员"]

JIJIAN_OFFICES = [f"第{c}纪检监察室" for c in "一二三四五六"]
GENDER_OPTIONS = ["男", "女"]
RISK_OPTIONS   = ["低", "高"]

# ★ 是否 / 有无 分开处理；默认分别为“否”“无”
SHIFO_OPTIONS = ["是", "否", "暂未掌握"]
YOUWU_OPTIONS = ["有", "无", "暂未掌握"]

# ★ 新增：婚姻/敏感度/关注度选项
MARRIAGE_OPTIONS = ["已婚", "未婚"]
SENSITIVITY_OPTIONS = ["普通", "敏感"]
ATTENTION_OPTIONS = ["普通", "高"]

# ★ 综合评估风险选项保持不变
ZONGHE_RISK_OPTIONS = ["低", "较低", "较高", "高"]

# 常规模板使用的字段分组（B05 字段会被剔除到单独页，但仍在同一向导里）
GROUP_DEFS = {
    "基本信息": ["第_纪检监察室", "填表日期"],
    "对象信息": ["对象姓名", "对象职务", "对象性别", "对象身份证号码", "谈话风险", "综合评估风险", "婚姻状况", "案件敏感程度", "案件社会关注度"],
    "时间地点": ["安全首课时间", "安全预案时间", "谈话方案时间",
               "谈话时间", "回访时间", "安全首课地点", "谈话地点"],
    "纪委人员": ["直接责任人", "第一责任人", "首谈领导",
               "主谈人", "参与人1","参与人2", "安全员", "安全员电话", "批准自行往返领导"],
    # 模板若有“主要任务”“谈话流程”等，也会自动出现在“其他字段”
}

# ---------- 文件名清洗 & 输出文件夹日期 ----------
def sanitize_filename(name: str) -> str:
    """去除 Windows 不合法字符"""
    return re.sub(r'[\\/:*?"<>|]', "_", name or "")

def _parse_to_date(s: str | None):
    if not s:
        return None
    s = str(s).strip()
    # 常见格式
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d", "%Y年%m月%d日"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except Exception:
            pass
    # 宽松匹配：YYYY 任意分隔 MM 任意分隔 DD
    m = re.search(r"(\d{4})\D{0,3}(\d{1,2})\D{0,3}(\d{1,2})", s)
    if m:
        y, mth, d = map(int, m.groups())
        try:
            return datetime.date(y, mth, d)
        except Exception:
            return None
    return None

def extract_folder_date(mapping: dict) -> str:
    """输出文件夹日期优先级：谈话时间 > 填表日期 > 今天；格式 YYYYMMDD"""
    for key in ("谈话时间", "填表日期"):
        d = _parse_to_date(mapping.get(key))
        if d:
            return d.strftime("%Y%m%d")
    return datetime.date.today().strftime("%Y%m%d")

# ================== 1. Word 替换工具 ==================
def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def find_placeholders(doc):
    ph = set()
    for p in iter_paragraphs(doc):
        ph.update(PLACEHOLDER_RE.findall("".join(r.text for r in p.runs)))
    return ph

def replace_in_paragraph(p, mapping):
    txt_org = "".join(r.text for r in p.runs)
    txt = txt_org
    for k, v in mapping.items():
        txt = txt.replace(f"`{k}`", str(v))
    if txt != txt_org:
        runs = p.runs or [p.add_run("")]
        runs[0].text = txt
        for r in runs[1:]:
            r.text = ""

def replace_everywhere(doc, mapping):
    for p in iter_paragraphs(doc):
        replace_in_paragraph(p, mapping)

def replace_role_phrase(doc, role):
    for p in iter_paragraphs(doc):
        txt_org = "".join(r.text for r in p.runs)
        txt = txt_org.replace(ROLE_PHRASE, role)
        if txt != txt_org:
            runs = p.runs or [p.add_run("")]
            runs[0].text = txt
            for r in runs[1:]:
                r.text = ""

def normalize_date(text):
    if not isinstance(text, str):
        try:
            text = text.strftime("%Y-%m-%d")
        except Exception:
            return str(text)
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d"):
        try:
            d = datetime.datetime.strptime(text.strip(), fmt).date()
            return f"{d.year}年{d.month:02d}月{d.day:02d}日"
        except ValueError:
            continue
    return text

# ★ 参与人2 统一格式化（空/“无”=>空串；否则前加“、”）
def format_canyuren2(val):
    s = (val or "").strip()
    s_plain = re.sub(r"[。\s]", "", s)
    if s_plain == "" or s_plain == "无":
        return ""
    return s if s.startswith("、") else ("、" + s)

# ================== 2. Excel 记录工具 ==================
def adjust_column_widths(xls_path: pathlib.Path):
    wb = openpyxl.load_workbook(xls_path)
    ws = wb.active
    for col in ws.columns:
        first_cell = next(iter(col))
        letter = get_column_letter(first_cell.column)
        width  = max(len(str(c.value)) if c.value else 0 for c in col) + 2
        ws.column_dimensions[letter].width = width
    wb.save(xls_path)
    wb.close()  # 显式关闭

def append_to_office_excel(record):
    """
    将记录追加到按科室命名的 Excel 中（<第x纪检监察室>.xlsx）。
    修复：只用一次 ExcelWriter 完成写入 + 调列宽，写完即关闭，再重命名。
    这样不会再由本进程占用临时文件。
    """
    office = (record.get("第_纪检监察室", "") or "").strip() or "未指定科室"
    path   = OUTPUT_DIR / f"{office}.xlsx"
    df_new = pd.DataFrame([record])

    # 读旧表（放在 with open(...) 里，确保读句柄立刻释放）
    if path.exists():
        try:
            with open(path, "rb") as f:
                df_old = pd.read_excel(f, dtype=str, engine="openpyxl")
        except Exception:
            df_old = pd.DataFrame()
        # 列对齐
        for col in df_new.columns.difference(df_old.columns):
            df_old[col] = ""
        for col in df_old.columns.difference(df_new.columns):
            df_new[col] = ""
        df = pd.concat([df_old, df_new], ignore_index=True)
    else:
        df = df_new

    # 写到临时文件
    tmp_fd, tmp_name = tempfile.mkstemp(
        suffix=".xlsx", prefix=path.stem + "_", dir=path.parent
    )
    os.close(tmp_fd)  # 立即关闭底层 fd（只保留路径）
    tmp_path = pathlib.Path(tmp_name)

    try:
        # 关键：用一次 ExcelWriter 完成全部写入 & 列宽设置，然后自动关闭释放句柄
        with pd.ExcelWriter(tmp_path, engine="openpyxl", mode="w") as writer:
            sheet_name = "Sheet1"
            df.to_excel(writer, index=False, sheet_name=sheet_name)

            # 用 writer.book / writer.sheets 调列宽（无需二次 open）
            wb = writer.book
            ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active

            for i, col in enumerate(df.columns, start=1):
                # 计算该列最大显示宽度（含表头）
                values = df[col].astype(str).replace("nan", "").tolist()
                max_len = max([len(str(col))] + [len(v) for v in values]) + 2
                ws.column_dimensions[get_column_letter(i)].width = max_len

            # with 结束会自动保存并关闭，一次到位

        # 句柄已释放，此时再做原子替换；失败时小退避重试几次
        for _ in range(8):
            try:
                os.replace(tmp_path, path)
                break
            except PermissionError:
                time.sleep(0.25)
        else:
            # 仍然失败：落盘到时间戳备份，提示用户
            ts = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            alt = path.with_name(f"{path.stem}-{ts}.xlsx")
            os.replace(tmp_path, alt)
            messagebox.showwarning(
                "目标文件被占用",
                f"目标文件可能被外部程序占用，已保存到：\n{alt}\n\n"
                f"请关闭占用的程序（如 Excel 预览/同步盘），再手动改名覆盖为：{path.name}",
                parent=menu if 'menu' in globals() else None
            )
    finally:
        # 清理孤儿临时文件
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception:
                pass


# ================== 3. 生成 Word 文档 ==================
def generate_docs(mapping):
    obj_raw = mapping.get("对象姓名", "未命名") or "未命名"
    obj = sanitize_filename(obj_raw)
    date_tag = extract_folder_date(mapping)  # ★ 输出文件夹日期：YYYYMMDD
    out   = OUTPUT_DIR / f"{obj}-{date_tag}"    # ★ 改为“对象姓名-YYYYMMDD”
    out.mkdir(parents=True, exist_ok=True)
    files = []

    for tpl in good_templates:
        if tpl.stem == SPECIAL_TEMPLATE_STEM:
            for role in ROLES:
                doc = Document(tpl)
                replace_everywhere(doc, mapping)
                replace_role_phrase(doc, role)
                name = tpl.stem
                for k, v in mapping.items():
                    name = name.replace(f"`{k}`", str(v))
                path = out / f"{name}-{role}.docx"
                doc.save(path); files.append(path)
        else:
            doc = Document(tpl)
            replace_everywhere(doc, mapping)
            name = tpl.stem
            for k, v in mapping.items():
                name = name.replace(f"`{k}`", str(v))
            path = out / f"{name}.docx"
            doc.save(path); files.append(path)
    return files

def process_record(mapping, collector: set, write_excel=True):
    for k in list(mapping.keys()):
        if ("日期" in k) or ("时间" in k):
            mapping[k] = normalize_date(mapping[k])
    if "参与人2" in mapping:
        mapping["参与人2"] = format_canyuren2(mapping.get("参与人2"))
    if write_excel:
        append_to_office_excel(mapping)
    collector.update(generate_docs(mapping))

# ================== 4. 扫描模板 & 占位符 ==================
good_templates, bad = [], []
for p in TEMPLATE_DIR.rglob("*.docx"):
    try:
        Document(p)
        good_templates.append(p)
    except (BadZipFile, PackageNotFoundError):
        bad.append(p)
if not good_templates:
    messagebox.showerror("未找到模板", "未发现可解析的 .docx 文件。")
    raise SystemExit

# 全部占位符、B05 占位符、非 B05 占位符
all_ph = set()
tpl_ph_map = {}
for tpl in good_templates:
    phs = find_placeholders(Document(tpl))
    tpl_ph_map[tpl.stem] = phs
    all_ph.update(phs)

b05_ph = tpl_ph_map.get(B05_TEMPLATE_STEM, set())
non_b05_ph = all_ph - b05_ph  # 常规模板填写内容

# 常规模板的分组键（不含 B05）
defaults = {
    "填表日期": datetime.date.today().strftime("%Y-%m-%d"),
    # ★ 默认“其他需要说明的情况”为 无（主要用于 B05）
    "其他需要说明的情况": "无",
}
group_keys_main = {g: [] for g in GROUP_DEFS}; group_keys_main["其他字段"] = []
for k in non_b05_ph:
    for g, lst in GROUP_DEFS.items():
        if k in lst:
            group_keys_main[g].append(k)
            break
    else:
        group_keys_main["其他字段"].append(k)

# ================== 5. 统一向导：常规 + B05（B05 两页） ==================
def wizard_fill_all():
    """一个窗口里完成常规字段和 B05（B05 拆成两页），仅 上一步/下一步/保存。
    保存按钮：只有在【最后一页】且所有必填项都填写后才点亮。
    说明：唯一可留空的字段是“参与人2”。
    """
    def is_shifo_key(key: str) -> bool:
        return "是否" in key
    def is_youwu_key(key: str) -> bool:
        return "有无" in key
    def is_date(k): return ("时间" in k) or ("日期" in k)

    mapping, page_vars, pages, page_keys = {}, {}, [], []
    wiz = tk.Toplevel()
    wiz.title("填写走读式谈话信息（含安全评估表）")
    wiz.resizable(False, True)

    # ---------- 生成常规模板分组页 ----------
    for grp, keys in group_keys_main.items():
        if not keys:
            continue
        frm = ttk.Frame(wiz, padding=18)
        ttk.Label(frm, text=grp, font=("微软雅黑", 12, "bold")
                 ).grid(columnspan=3, pady=(0, 10))
        for r, key in enumerate(sorted(keys), 1):
            ttk.Label(frm, text=key + "：").grid(row=r, column=0,
                                                sticky="e", padx=6, pady=4)
            # 下拉/输入控件
            if key == "第_纪检监察室":
                var = tk.StringVar(value=JIJIAN_OFFICES[0])
                ent = ttk.Combobox(frm, width=28, state="readonly",
                                   values=JIJIAN_OFFICES, textvariable=var)
            elif key == "对象性别":
                var = tk.StringVar(value=GENDER_OPTIONS[0])
                ent = ttk.Combobox(frm, width=8, state="readonly",
                                   values=GENDER_OPTIONS, textvariable=var)
            elif key == "谈话风险":
                var = tk.StringVar(value=RISK_OPTIONS[0])
                ent = ttk.Combobox(frm, width=8, state="readonly",
                                   values=RISK_OPTIONS, textvariable=var)
            elif key == "综合评估风险":
                var = tk.StringVar(value=ZONGHE_RISK_OPTIONS[0])
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=ZONGHE_RISK_OPTIONS, textvariable=var)
            elif key == "婚姻状况":
                var = tk.StringVar(value=MARRIAGE_OPTIONS[0])
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=MARRIAGE_OPTIONS, textvariable=var)
            elif key == "案件敏感程度":
                var = tk.StringVar(value=SENSITIVITY_OPTIONS[0])
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=SENSITIVITY_OPTIONS, textvariable=var)
            elif key == "案件社会关注度":
                var = tk.StringVar(value=ATTENTION_OPTIONS[0])
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=ATTENTION_OPTIONS, textvariable=var)
            elif is_shifo_key(key):
                var = tk.StringVar(value=SHIFO_OPTIONS[1])  # 默认 否
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=SHIFO_OPTIONS, textvariable=var)
            elif is_youwu_key(key):
                var = tk.StringVar(value=YOUWU_OPTIONS[1])  # 默认 无
                ent = ttk.Combobox(frm, width=10, state="readonly",
                                   values=YOUWU_OPTIONS, textvariable=var)
            elif key == "参与人2":
                d = defaults.get(key, "")
                var = tk.StringVar(value=d)
                ent = ttk.Entry(frm, width=30, textvariable=var)
                ttk.Label(frm, text="（如无请留空）").grid(row=r, column=2, sticky="w", padx=2)
            else:
                d = defaults.get(key, "")
                if not d and is_date(key):
                    d = datetime.date.today().strftime("%Y-%m-%d")
                width = 60 if grp == "其他字段" else 30
                var = tk.StringVar(value=d)
                ent = ttk.Entry(frm, width=width, textvariable=var)
                if key == "主要任务":
                    ttk.Label(frm, text="示例：与张三核实XXX内容").grid(row=r, column=2, sticky="w", padx=2)
                if key == "谈话流程":
                    ttk.Label(frm, text="示例：先详细听取张三XXX内容，后就XXX内容进行提问，了解XXX"
                              ).grid(row=r, column=2, sticky="w", padx=2)

            ent.grid(row=r, column=1, sticky="w", padx=6, pady=4)
            page_vars[key] = var
        pages.append(frm)
        page_keys.append(keys)

    # ---------- 生成 B05 两页（同一向导里继续） ----------
    if b05_ph:
        # 排序：将“其他需要说明的情况”放到最后
        ordered_keys = sorted([k for k in b05_ph if k != "其他需要说明的情况"])
        if "其他需要说明的情况" in b05_ph:
            ordered_keys.append("其他需要说明的情况")
        half = (len(ordered_keys) + 1) // 2
        parts = [ordered_keys[:half], ordered_keys[half:]]

        for idx, keys in enumerate(parts, start=1):
            frm = ttk.Frame(wiz, padding=18)
            ttk.Label(frm, text=f"安全评估表（{idx}/2）",
                      font=("微软雅黑", 12, "bold")).grid(columnspan=3, pady=(0, 10))
            row = 1
            for key in keys:
                ttk.Label(frm, text=key + "：").grid(row=row, column=0, sticky="e", padx=6, pady=4)

                # 默认值：日期给今天；“其他需要说明的情况”给“无”
                d = defaults.get(key, "")
                if not d and is_date(key):
                    d = datetime.date.today().strftime("%Y-%m-%d")
                if key == "其他需要说明的情况":
                    d = "无"

                # 下拉逻辑与常规模板一致
                if key == "婚姻状况":
                    var = tk.StringVar(value=MARRIAGE_OPTIONS[0])
                    ent = ttk.Combobox(frm, width=12, state="readonly",
                                       values=MARRIAGE_OPTIONS, textvariable=var)
                elif key == "案件敏感程度":
                    var = tk.StringVar(value=SENSITIVITY_OPTIONS[0])
                    ent = ttk.Combobox(frm, width=10, state="readonly",
                                       values=SENSITIVITY_OPTIONS, textvariable=var)
                elif key == "案件社会关注度":
                    var = tk.StringVar(value=ATTENTION_OPTIONS[0])
                    ent = ttk.Combobox(frm, width=10, state="readonly",
                                       values=ATTENTION_OPTIONS, textvariable=var)
                elif is_shifo_key(key):
                    var = tk.StringVar(value=SHIFO_OPTIONS[1])  # 默认 否
                    ent = ttk.Combobox(frm, width=10, state="readonly",
                                       values=SHIFO_OPTIONS, textvariable=var)
                elif is_youwu_key(key):
                    var = tk.StringVar(value=YOUWU_OPTIONS[1])  # 默认 无
                    ent = ttk.Combobox(frm, width=10, state="readonly",
                                       values=YOUWU_OPTIONS, textvariable=var)
                else:
                    var = tk.StringVar(value=d)
                    ent = ttk.Entry(frm, width=50, textvariable=var)
                    if key == "主要任务":
                        ttk.Label(frm, text="示例：与张三核实XXX内容"
                                 ).grid(row=row, column=2, sticky="w", padx=2)
                    if key == "谈话流程":
                        ttk.Label(frm, text="示例：先详细听取张三XXX内容，后就XXX内容进行提问，了解XXX"
                                 ).grid(row=row, column=2, sticky="w", padx=2)

                ent.grid(row=row, column=1, sticky="w", padx=6, pady=4)
                page_vars[key] = var
                row += 1

            pages.append(frm)
            page_keys.append(keys)

    # ---------- 导航区：仅 上一步 / 下一步 / 保存 ----------
    nav = ttk.Frame(wiz, padding=(12, 6))
    nav.pack(side="bottom", fill="x")
    b_prev = ttk.Button(nav, text="← 上一步")
    b_next = ttk.Button(nav, text="下一步 →")
    b_save = ttk.Button(nav, text="保存", state="disabled")  # 初始灰色
    b_prev.grid(row=0, column=0, padx=4)
    b_next.grid(row=0, column=1, padx=4)
    b_save.grid(row=0, column=2, padx=4)

    cur = 0

    # 校验：除“参与人2”外全部必填
    def page_valid(i):
        return all(page_vars[k].get().strip() for k in page_keys[i] if k != "参与人2")

    def all_valid():
        return all(v.get().strip() for k, v in page_vars.items() if k != "参与人2")

    def update_save_state(*_):
        # 只有在最后一页且全部填写完成才点亮保存
        last_page = (cur == len(pages) - 1)
        b_save['state'] = 'normal' if (last_page and all_valid()) else 'disabled'

    def show(i):
        nonlocal cur
        pages[cur].pack_forget()
        cur = i
        pages[cur].pack()
        b_prev['state'] = 'normal' if cur > 0 else 'disabled'
        b_next['state'] = 'normal' if cur < len(pages) - 1 else 'disabled'
        update_save_state()

    b_prev.config(command=lambda: show(cur - 1))

    def go_next(e=None):
        if not page_valid(cur):
            messagebox.showwarning("请完善", "当前页还有未填写字段！", parent=wiz)
            return
        if cur < len(pages) - 1:
            show(cur + 1)
    b_next.config(command=go_next)
    wiz.bind("<Return>", go_next)

    finished = False
    def on_save():
        if not all_valid() or cur != len(pages) - 1:
            # 双保险，正常情况下按钮是灰的
            messagebox.showwarning("请完善", "还有未填写字段，或未到最后一页！", parent=wiz)
            return
        nonlocal finished
        finished = True
        wiz.destroy()
    b_save.config(command=on_save)

    # 变量变动时联动更新“保存”按钮状态
    for var in page_vars.values():
        try:
            var.trace_add('write', update_save_state)
        except Exception:
            # 兼容旧版 Tk
            try:
                var.trace('w', update_save_state)
            except Exception:
                pass

    # 关闭即放弃
    wiz.bind("<Escape>", lambda e: wiz.destroy())
    wiz.protocol("WM_DELETE_WINDOW", wiz.destroy)

    # 展示第一页
    pages[0].pack()
    b_prev["state"] = "disabled"
    update_save_state()
    wiz.grab_set()
    wiz.wait_window()

    if not finished:
        return None
    mapping.update({k: v.get().strip() for k, v in page_vars.items()})
    return mapping

# ================== 6. Excel 导入 ==================
def import_from_excel():
    f = filedialog.askopenfilename(filetypes=[("Excel 文件", "*.xlsx *.xls")])
    if not f:
        return
    try:
        df = pd.read_excel(f, dtype=str, engine="openpyxl")
    except Exception as e:
        messagebox.showerror("读取失败", f"无法读取 Excel：\n{e}", parent=menu)
        return
    if "对象姓名" not in df.columns:
        messagebox.showwarning("列缺失", "Excel 必须包含 “对象姓名” 列！", parent=menu)
        return
    df.fillna("", inplace=True)

    sel = tk.Toplevel()
    sel.title("选择生成对象")
    sel.resizable(False, False)
    ttk.Label(sel, text="勾选需要生成的对象：",
              font=("微软雅黑", 10, "bold")).pack(pady=(10, 6))
    lb = tk.Listbox(sel, selectmode="extended", width=40, height=15)
    for i, n in enumerate(df["对象姓名"]):
        lb.insert(i, n)
    lb.pack(padx=12)
    frm = ttk.Frame(sel, padding=6)
    frm.pack()
    ttk.Button(frm, text="全选",
               command=lambda: lb.selection_set(0, tk.END)).grid(row=0, column=0, padx=4)
    gen = ttk.Button(frm, text="生成")
    gen.grid(row=0, column=1, padx=4)

    outs: set[pathlib.Path] = set()
    def doit():
        idxs = lb.curselection()
        if not idxs:
            messagebox.showinfo("未选择", "请至少选择一个对象！", parent=sel)
            return
        sel.destroy()
        for i in idxs:
            rec = df.iloc[i].to_dict()
            # 填充缺失列（包含 B05 在内的所有占位符）
            for k in all_ph:
                # ★ 为“其他需要说明的情况”提供默认值“无”
                if k == "其他需要说明的情况" and not rec.get(k, "").strip():
                    rec[k] = "无"
                rec.setdefault(k, "")
            process_record(rec, outs, write_excel=False)  # 自动规范参与人2/日期
        paths = sorted(map(str, outs))
        messagebox.showinfo("完成", "已生成以下文件：\n" + "\n".join(paths), parent=menu)
    gen.config(command=doit)
    sel.grab_set()

# ================== 7. 主菜单 ==================
root = tk.Tk()
root.withdraw()
menu = tk.Toplevel()
menu.title("温岭纪委走读式谈话文件生成工具")
menu.geometry("520x320")
menu.resizable(False, False)

ttk.Label(menu, text="温岭纪委走读式谈话文件生成工具",
          font=("微软雅黑", 16, "bold")).pack(pady=(20, 4))
ttk.Label(menu, text="© 温岭纪委六室 单柳昊",
          font=("微软雅黑", 9)).pack()

btn_frame = ttk.Frame(menu, padding=20)
btn_frame.pack(expand=True)
b_manual = ttk.Button(btn_frame, text="手动填写", width=26)
b_excel  = ttk.Button(btn_frame, text="从 Excel 导入", width=26)
b_manual.grid(row=0, column=0, pady=6)
b_excel.grid(row=1, column=0, pady=6)
ttk.Label(menu, text="生成结果仅供参考，请仔细检查",
          font=("微软雅黑", 8)).place(relx=1.0, rely=1.0,
                                      anchor="se", x=-6, y=-4)

def run_manual():
    # 统一向导（常规 + B05 两页）
    mapping = wizard_fill_all()
    if mapping is None:
        return

    outs: set[pathlib.Path] = set()
    process_record(mapping, outs)
    paths = sorted(map(str, outs))
    messagebox.showinfo("完成", "已生成以下文件：\n" + "\n".join(paths), parent=menu)

b_manual.config(command=run_manual)
b_excel.config(command=import_from_excel)

# ★ 统一退出，确保不留后台进程
def exit_app(*_):
    try:
        for w in list(root.winfo_children()):
            try:
                w.destroy()
            except Exception:
                pass
        root.quit()
    finally:
        try:
            root.destroy()
        finally:
            sys.exit(0)

menu.bind("<Escape>", exit_app)
menu.protocol("WM_DELETE_WINDOW", exit_app)

menu.mainloop()
